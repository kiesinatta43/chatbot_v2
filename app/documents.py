from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

try:
    import fitz as _fitz  # pymupdf
    _HAS_FITZ = True
except ImportError:
    _HAS_FITZ = False
    from pypdf import PdfReader

from app.config import CHUNK_OVERLAP, CHUNK_SIZE, SRC_DIR

try:
    from docx import Document as DocxDocument
except ModuleNotFoundError:
    DocxDocument = None


@dataclass
class DocumentChunk:
    source: str
    section: str
    text: str


@dataclass
class SourceDocument:
    source: str
    fingerprint: str
    text: str


def load_supported_documents() -> list[DocumentChunk]:
    chunks: list[DocumentChunk] = []
    for document in load_source_documents():
        chunks.extend(split_into_chunks(document.source, document.text))
    return [c for c in chunks if not is_noise_chunk(c.text)]


def is_noise_chunk(text: str) -> bool:
    normalized = text.strip()
    if len(normalized) < 40:
        return True
    noise_markers = [
        "สงวนลิขสิทธิ์", "บริษัทพัฒนาคุณภาพวิชาการ", "ActiveLearning",
        "GPAS 5 Steps", "LiveFacebook", "PhotoEditingChallenge",
        "แฟกซ์:", "โทร. ๐-", "สุดยอดคู่มือครู",
    ]
    for marker in noise_markers:
        if marker in normalized:
            return True
    thai_count = len(re.findall(r"[ก-๙]", normalized))
    return thai_count > 0 and thai_count / len(normalized) < 0.25


def load_source_documents() -> list[SourceDocument]:
    documents: list[SourceDocument] = []
    for path in sorted(SRC_DIR.rglob("*")):
        if not path.is_file():
            continue
        text = read_file(path)
        if not text.strip():
            continue
        source = path.relative_to(SRC_DIR).as_posix()
        documents.append(
            SourceDocument(
                source=source,
                fingerprint=build_file_fingerprint(path, source),
                text=text,
            )
        )
    return documents


def build_source_manifest() -> dict[str, str]:
    manifest: dict[str, str] = {}
    for path in sorted(SRC_DIR.rglob("*")):
        if not path.is_file():
            continue
        source = path.relative_to(SRC_DIR).as_posix()
        manifest[source] = build_file_fingerprint(path, source)
    return manifest


def _extract_pdf_with_fitz(path: str) -> str:
    doc = _fitz.open(path)
    pages: list[str] = []
    for page in doc:
        # "dict" mode gives word-level bounding boxes so we can reconstruct
        # reading order and skip watermark text (very large font or low opacity)
        blocks = page.get_text("dict", flags=_fitz.TEXT_PRESERVE_LIGATURES)["blocks"]
        lines: list[str] = []
        for block in blocks:
            if block.get("type") != 0:  # skip image blocks
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                # skip spans that look like watermarks: very large font or colour close to white
                filtered = []
                for span in spans:
                    size = span.get("size", 0)
                    color = span.get("color", 0)
                    r = (color >> 16) & 0xFF
                    g = (color >> 8) & 0xFF
                    b = color & 0xFF
                    is_light = r > 180 and g > 180 and b > 180
                    is_huge = size > 40
                    if not is_light and not is_huge:
                        filtered.append(span.get("text", ""))
                text = "".join(filtered).strip()
                if text:
                    lines.append(text)
        if lines:
            pages.append("\n".join(lines))
    doc.close()
    return "\n\n".join(pages)


def read_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return normalize_document_text(path.read_text(encoding="utf-8", errors="ignore"))
    if suffix == ".pdf":
        if _HAS_FITZ:
            return normalize_document_text(_extract_pdf_with_fitz(str(path)))
        reader = PdfReader(str(path))
        return normalize_document_text("\n".join(page.extract_text() or "" for page in reader.pages))
    if suffix == ".docx":
        if DocxDocument is None:
            raise ModuleNotFoundError(
                "python-docx is required to ingest .docx files. Install dependencies from requirements.txt "
                "or remove .docx files from src/."
            )
        doc = DocxDocument(str(path))
        return normalize_document_text("\n".join(paragraph.text for paragraph in doc.paragraphs))
    return ""


def split_into_chunks(source: str, text: str) -> list[DocumentChunk]:
    normalized = normalize_for_chunking(text)
    if len(normalized) <= CHUNK_SIZE:
        return [DocumentChunk(source=source, section=detect_section(normalized, "ส่วนที่ 1"), text=normalized)]

    segments = split_into_segments(normalized)
    chunks: list[DocumentChunk] = []
    index = 1
    current = ""

    for segment in segments:
        if not current:
            current = segment
            continue

        candidate = f"{current} {segment}".strip()
        if len(candidate) <= chunk_limit_for_segment(current, segment):
            current = candidate
            continue

        chunks.append(DocumentChunk(source=source, section=detect_section(current, f"ส่วนที่ {index}"), text=current))
        index += 1
        overlap = current[-CHUNK_OVERLAP:].strip()
        if overlap:
            first_space = overlap.find(" ")
            if 0 < first_space < 40:
                overlap = overlap[first_space:].strip()
        current = f"{overlap} {segment}".strip() if overlap else segment

    if current:
        chunks.append(DocumentChunk(source=source, section=detect_section(current, f"ส่วนที่ {index}"), text=current))

    return chunks


def detect_section(text: str, fallback: str) -> str:
    patterns = [
        r"(หน่วยที่\s*\d+[^\n]{0,30})",
        r"(บทที่\s*\d+[^\n]{0,30})",
        r"(\d+\.\d+\s+[ก-๙]{4,}[^\n]{0,25})",
        r"(กระบวนการออกแบบเชิงวิศวกรรม)",
        r"(เทคนิค 5W1H)",
        r"(ขั้น(?:ระบุปัญหา|รวบรวมข้อมูล|ออกแบบวิธีการ|วางแผน|ทดสอบ|นำเสนอ)[^\n]{0,20})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(1).strip()[:50]
    return fallback


def split_into_segments(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?])\s+|\s*(?<=।)\s*|\n+", text)
    segments = [part.strip() for part in parts if part.strip()]
    return segments or [text]


def build_file_fingerprint(path: Path, source: str) -> str:
    stats = path.stat()
    return f"{source}:{stats.st_size}:{int(stats.st_mtime)}"


def normalize_document_text(text: str) -> str:
    lines = [clean_extracted_line(line) for line in text.splitlines()]
    cleaned_lines = [line for line in lines if line]

    merged: list[str] = []
    for line in cleaned_lines:
        if merged and should_merge_lines(merged[-1], line):
            merged[-1] = merge_line_pair(merged[-1], line)
        else:
            merged.append(line)

    text = "\n".join(merged)
    text = normalize_ocr_artifacts(text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def normalize_for_chunking(text: str) -> str:
    normalized = text.replace("\n\n", " <PARA> ")
    normalized = " ".join(normalized.split())
    normalized = normalized.replace(" <PARA> ", "\n\n")
    return normalized.strip()


def clean_extracted_line(line: str) -> str:
    line = line.replace("\u00ad", "")
    line = line.replace("\ufeff", "")
    line = line.replace("\ufffd", "")
    line = re.sub(r"\s+", " ", line).strip()
    line = re.sub(r"(?<=[A-Za-z])\s+(?=[A-Za-z])", "", line)
    line = merge_broken_thai_spacing(line)
    line = re.sub(r"\s+([.,;:!?])", r"\1", line)
    line = normalize_ocr_artifacts(line)
    return line


def merge_broken_thai_spacing(text: str) -> str:
    return re.sub(r"(?<=[ก-๙])\s+(?=[ก-๙])", "", text)


def should_merge_lines(previous: str, current: str) -> bool:
    if not previous or not current:
        return False
    if previous.endswith((".", "?", "!", ":", ";")):
        return False
    if re.match(r"^(บทที่|หน่วยที่|กิจกรรม|ใบงาน|แบบฝึกหัด|หัวข้อ|[0-9]+\.)", current):
        return False
    return True


def merge_line_pair(previous: str, current: str) -> str:
    if previous.endswith("-"):
        return previous[:-1] + current
    if previous.endswith(("(", "/", "“", "\"")):
        return previous + current
    return f"{previous} {current}"


def chunk_limit_for_segment(current: str, next_segment: str) -> int:
    if looks_like_enumeration(current) or looks_like_enumeration(next_segment):
        return CHUNK_SIZE + 500
    return CHUNK_SIZE


def looks_like_enumeration(text: str) -> bool:
    return bool(
        re.search(
            r"(ได้แก่|ประกอบด้วย|มี\s*\d+\s*ขั้นตอน|ขั้นระบุปัญหา|ขั้นรวบรวมข้อมูล|ขั้นออกแบบวิธีการแก้ปัญหา|ขั้นวางแผน|ขั้นทดสอบ|ขั้นนำเสนอ)",
            text,
        )
    )


def normalize_ocr_artifacts(text: str) -> str:
    cleaned = text
    replacements = {
        "ความี": "ความ",
        "เปั็น": "เป็น",
        "ขั้อง": "ของ",
        "คิวิาม": "ความ",
        "ขั้ั": "ขั้น",
        "ขั้�น": "ขั้น",
        "ขั้้น": "ขั้น",
        "ตััวอย่่าง": "ตัวอย่าง",
        "คำตัอบ": "คำตอบ",
        "สำาคัญ": "สำคัญ",
        "ทำางาน": "ทำงาน",
        "ดำาเนิน": "ดำเนิน",
        "ปััญหา": "ปัญหา",
        "ก่าร": "การ",
        "ข้้อมูล": "ข้อมูล",
        "เก่ิด": "เกิด",
        "อย่่าง": "อย่าง",
        "ชิ้่วย่": "ช่วย",
        "ด้้วย่": "ด้วย",
        "เป่็น": "เป็น",
        "ตัรง": "ตรง",
        "ที่ำ": "ทำ",
        "สิงที่ี": "สิ่งที่",
        "มีอง": "มอง",
        "ย่าก่": "ยาก",
        "ได้้ง่าย่": "ได้ง่าย",
        "เพื่ือ": "เพื่อ",
        "สือสาร": "สื่อสาร",
        "อืน": "อื่น",
        "เขั้้าใจ": "เข้าใจ",
        "หมีาย่": "หมาย",
        "ใชิ้้": "ใช้",
        "แนวที่าง": "แนวทาง",
        "เพื่ือนำไปัสู่": "เพื่อนำไปสู่",
        "ปัฏิิบัตัิงาน": "ปฏิบัติงาน",
        "ตั่อไปั": "ต่อไป",
        "ซอฟ้ต์แวร์": "ซอฟต์แวร์",
        "ตั้นแบบ": "ต้นแบบ",
        "แก่้ปััญหา": "แก้ปัญหา",
        "แก่้ปัญหา": "แก้ปัญหา",
        "ระบุปััญหา": "ระบุปัญหา",
        "วิเครูาะห์": "วิเคราะห์",
        "ขั้้อมีูล": "ข้อมูล",
        "รูวบรูวมู": "รวบรวม",
        "ไมู่มูากพอ": "ไม่มากพอ",
        "จ่ะทำให้": "จะทำให้",
        "ด้ำเนินการู": "ดำเนินการ",
        "ผูู้้ใช้้": "ผู้ใช้",
        "รูถย่นตั์": "รถยนต์",
        "ครูัง": "ครั้ง",
        "หรูือ": "หรือ",
        "ข้ณะ": "ขณะ",
    }
    for original, replacement in replacements.items():
        cleaned = cleaned.replace(original, replacement)

    cleaned = re.sub(r"[�′]+", "", cleaned)
    cleaned = re.sub(r"([ัิ-ฺ็-๎])\1+", r"\1", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()